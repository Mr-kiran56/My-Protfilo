from langchain_text_splitters import TokenTextSplitter
from langchain_community.docstore.document import Document
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument
import os


def file_process(file_path: str):


    print(f"\nProcessing file: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    text_overall = ""

    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        print(f"Loaded {len(pages)} pages from PDF")
        for page in pages:
            text_overall += page.page_content + "\n"

    elif ext in [".docx", ".doc"]:
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            text_overall += para.text + "\n"
        print(f"Loaded DOCX with {len(doc.paragraphs)} paragraphs")

    elif ext in [".txt", ".md"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text_overall = f.read()
        print(f"Loaded TXT/MD with {len(text_overall.splitlines())} lines")

    else:
        raise ValueError(f"Unsupported file type: {ext}")


    big_splitter = TokenTextSplitter(
        chunk_size=1000,     
        chunk_overlap=100
    )
    big_chunks = big_splitter.split_text(text_overall)
    print(f"Created {len(big_chunks)} large chunks")

    documents = [Document(page_content=t) for t in big_chunks]

    
    small_splitter = TokenTextSplitter(
        chunk_size=300,      
        chunk_overlap=50     
    )
    final_chunks = small_splitter.split_documents(documents)
    print(f"Created {len(final_chunks)} final chunks")

    return final_chunks
